"""Generate the Staff Biodata Management System user training guide as a Word document."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_background(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_tip_box(doc, title, body_lines):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_background(cell, 'FFF4CE')
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    cell.width = Cm(15)

    p = cell.paragraphs[0]
    run = p.add_run(f"  {title}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x8A, 0x6D, 0x00)

    for line in body_lines:
        para = cell.add_paragraph()
        para.paragraph_format.left_indent = Cm(0.3)
        r = para.add_run(line)
        r.font.size = Pt(11)

    # Add border to table
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')
        border.set(qn('w:color'), 'E0B400')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    doc.add_paragraph()


def add_warning_box(doc, title, body_lines):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_background(cell, 'FDE7E9')

    p = cell.paragraphs[0]
    run = p.add_run(f"  {title}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xA8, 0x00, 0x00)

    for line in body_lines:
        para = cell.add_paragraph()
        para.paragraph_format.left_indent = Cm(0.3)
        r = para.add_run(line)
        r.font.size = Pt(11)

    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')
        border.set(qn('w:color'), 'C03030')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    doc.add_paragraph()


def add_screenshot_placeholder(doc, caption):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_background(cell, 'F0F0F0')

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n[ SCREENSHOT PLACEHOLDER ]\n")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    cap = cell.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption + "\n")
    cap_run.italic = True
    cap_run.font.size = Pt(10)
    cap_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'dashed')
        border.set(qn('w:sz'), '6')
        border.set(qn('w:color'), 'A0A0A0')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    doc.add_paragraph()


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        if level == 0:
            run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)
        elif level == 1:
            run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)
        else:
            run.font.color.rgb = RGBColor(0x2E, 0x5C, 0x9A)
    return h


def add_steps(doc, steps):
    for step in steps:
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(step)
        run.font.size = Pt(11)


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.bold = bold
    r.italic = italic
    return p


def add_field_row(table, name, description):
    row = table.add_row()
    row.cells[0].text = ""
    row.cells[1].text = ""

    p1 = row.cells[0].paragraphs[0]
    r1 = p1.add_run(name)
    r1.bold = True
    r1.font.size = Pt(10)

    p2 = row.cells[1].paragraphs[0]
    r2 = p2.add_run(description)
    r2.font.size = Pt(10)


def build_document():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # === COVER ===
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("\n\n\nSTAFF BIODATA MANAGEMENT SYSTEM")
    title_run.bold = True
    title_run.font.size = Pt(26)
    title_run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("\nUser Training Guide")
    sub_run.bold = True
    sub_run.font.size = Pt(18)
    sub_run.font.color.rgb = RGBColor(0x2E, 0x5C, 0x9A)

    org = doc.add_paragraph()
    org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org_run = org.add_run("\n\nFor the Administration Staff of the\nCustomary Court of Appeal")
    org_run.font.size = Pt(14)
    org_run.italic = True

    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ver_run = ver.add_run("\n\n\n\n\n\nVersion 1.0\nTraining Edition")
    ver_run.font.size = Pt(12)

    doc.add_page_break()

    # === WELCOME ===
    add_heading(doc, "Welcome", level=1)
    add_para(doc,
        "This guide will help you use the Staff Biodata Management System. "
        "You do not need to be a computer expert to follow it. Every task is "
        "explained step by step, in plain language."
    )
    add_para(doc,
        "Read the guide once from start to finish. After that, you can come "
        "back to any section when you need a reminder."
    )

    add_tip_box(doc, "TIP", [
        "Keep this guide near your computer.",
        "If something on the screen does not match this guide, call the "
        "system administrator before clicking anything."
    ])

    # === TABLE OF CONTENTS ===
    add_heading(doc, "What This Guide Covers", level=1)
    toc = [
        "1.  How to log in and log out safely",
        "2.  How to search for a staff member",
        "3.  How to view a staff profile",
        "4.  How to add a new staff record",
        "5.  How to update a staff record",
        "6.  How to record a promotion or transfer",
        "7.  How to view the promotions due list",
        "8.  How to view the retirement monitor",
        "9.  How to generate and download a report",
        "10. What to do if you forget your password",
        "11. Common errors and what they mean",
    ]
    for item in toc:
        add_para(doc, item)

    doc.add_page_break()

    # === 1. LOGIN ===
    add_heading(doc, "1. How to Log In and Log Out Safely", level=1)
    add_para(doc,
        "Logging in tells the system who you are. Logging out keeps your "
        "records safe when you leave your desk."
    )

    add_heading(doc, "To log in", level=2)
    add_steps(doc, [
        "Open your web browser (Chrome or Edge).",
        "Type the system address given to you by the administrator into the address bar at the top.",
        "Press the Enter key on your keyboard.",
        "The login page will appear.",
        "Click inside the box labelled \"Username\" and type your username.",
        "Click inside the box labelled \"Password\" and type your password.",
        "Click the blue \"Log In\" button.",
        "Wait a few seconds. The main dashboard will appear on your screen.",
    ])

    add_screenshot_placeholder(doc, "Figure 1.1 - The login page")

    add_heading(doc, "To log out", level=2)
    add_steps(doc, [
        "Look at the top-right corner of the screen.",
        "Click on your name or the small profile icon.",
        "A short menu will drop down.",
        "Click \"Log Out\".",
        "You will be taken back to the login page. You are now safely logged out.",
    ])

    add_screenshot_placeholder(doc, "Figure 1.2 - The Log Out option in the user menu")

    add_tip_box(doc, "SAFETY TIPS", [
        "Never share your password with anyone, even a colleague.",
        "Always log out before leaving your desk for a long time.",
        "Do not let the browser \"remember\" your password on a shared computer.",
        "If you suspect someone knows your password, change it the same day.",
    ])

    doc.add_page_break()

    # === 2. SEARCH ===
    add_heading(doc, "2. How to Search for a Staff Member", level=1)
    add_para(doc,
        "You can find any staff record quickly using the search box. You can "
        "search by name, file number, or department."
    )

    add_steps(doc, [
        "From the dashboard, click \"Staff Records\" on the left-hand menu.",
        "The staff list page will open.",
        "Find the long search box at the top of the page.",
        "Click inside the search box.",
        "Type the staff member's name, file number, or department.",
        "As you type, the list below will narrow down automatically.",
        "When you see the correct name, click on it to open the profile.",
    ])

    add_screenshot_placeholder(doc, "Figure 2.1 - The search box on the Staff Records page")

    add_tip_box(doc, "HELPFUL TIPS", [
        "You do not need to type the full name. \"Ade\" will find \"Adebayo\", \"Adekunle\", and so on.",
        "If you do not see the staff member, check the spelling.",
        "You can also click the \"Filter\" button to narrow by department, grade, or status.",
    ])

    doc.add_page_break()

    # === 3. VIEW PROFILE ===
    add_heading(doc, "3. How to View a Staff Profile", level=1)
    add_para(doc,
        "A staff profile shows everything the system knows about one person: "
        "personal details, employment history, promotions, and documents."
    )

    add_steps(doc, [
        "Find the staff member using the search (see Section 2).",
        "Click the staff member's name in the list.",
        "The profile page will open.",
        "You will see several tabs at the top: Personal, Employment, Promotions, Documents, History.",
        "Click any tab to see that section.",
        "To go back to the staff list, click \"Back to Staff List\" or the back arrow at the top-left.",
    ])

    add_screenshot_placeholder(doc, "Figure 3.1 - A typical staff profile showing the tabs")

    add_tip_box(doc, "TIP", [
        "Viewing a profile does NOT change anything. You can click around freely.",
        "Changes only happen when you click \"Save\" or \"Update\".",
    ])

    doc.add_page_break()

    # === 4. ADD NEW STAFF ===
    add_heading(doc, "4. How to Add a New Staff Record", level=1)
    add_para(doc,
        "Use this when a new officer joins the Court. Take your time and "
        "fill the form carefully."
    )

    add_steps(doc, [
        "Click \"Staff Records\" on the left-hand menu.",
        "Click the green \"+ Add New Staff\" button at the top-right.",
        "A blank form will open.",
        "Fill in each field, one at a time (see the list below).",
        "Fields marked with a red star (*) MUST be filled.",
        "When you finish, click the blue \"Save\" button at the bottom.",
        "A green message will say \"Staff record added successfully\".",
        "The new staff will now appear in the staff list.",
    ])

    add_screenshot_placeholder(doc, "Figure 4.1 - The Add New Staff form")

    add_heading(doc, "Each field explained", level=2)

    fields_table = doc.add_table(rows=1, cols=2)
    fields_table.style = 'Light Grid Accent 1'
    hdr = fields_table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "What to enter"
    for cell in hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    fields = [
        ("File Number *", "The unique staff file number given by HR. Example: CCA/2024/0123."),
        ("Surname *", "The staff member's last name (family name)."),
        ("First Name *", "The staff member's first name."),
        ("Middle Name", "The middle name, if any. Leave blank if none."),
        ("Date of Birth *", "Click the small calendar icon and pick the date. Format is day/month/year."),
        ("Gender *", "Click the drop-down and choose Male or Female."),
        ("Marital Status", "Single, Married, Widowed, or Divorced."),
        ("State of Origin *", "Choose from the drop-down list of states."),
        ("Local Government Area", "Choose the LGA after picking the state."),
        ("Home Address", "The staff member's residential address."),
        ("Phone Number *", "11-digit phone number, e.g. 08031234567."),
        ("Email Address", "A working email address, if available."),
        ("Date of First Appointment *", "The day the officer first joined the Court."),
        ("Current Grade Level *", "Choose the grade level, e.g. GL08, GL10."),
        ("Current Step", "Choose the step within the grade, e.g. Step 3."),
        ("Department / Unit *", "Choose the department where the officer works."),
        ("Designation / Post *", "The job title, e.g. \"Senior Court Clerk\"."),
        ("Cadre", "The job cadre, e.g. Administrative, Judicial Support, Secretarial."),
        ("Qualification(s)", "Highest academic qualification, e.g. BSc, HND, NCE."),
        ("Date of Last Promotion", "The date of the most recent promotion."),
        ("Date of Confirmation", "The date the officer was confirmed in service."),
        ("Date of Retirement", "Calculated automatically once date of birth and first appointment are filled."),
        ("Next of Kin Name", "Name of the person to contact in case of emergency."),
        ("Next of Kin Phone", "Phone number of the next of kin."),
        ("Photograph", "Click \"Upload Photo\" and choose a passport-style picture from your computer."),
    ]
    for name, desc in fields:
        add_field_row(fields_table, name, desc)

    doc.add_paragraph()

    add_tip_box(doc, "TIP", [
        "Double-check the file number and date of birth before saving. These are hard to correct later.",
        "If you are interrupted, the form will lose what you have typed. Finish in one sitting if you can.",
    ])

    doc.add_page_break()

    # === 5. UPDATE ===
    add_heading(doc, "5. How to Update a Staff Record", level=1)
    add_para(doc,
        "Use this when something changes - a new phone number, a new address, "
        "a new qualification, and so on."
    )

    add_steps(doc, [
        "Find the staff member using the search (see Section 2).",
        "Open the profile by clicking the name.",
        "Click the \"Edit\" button at the top-right of the profile.",
        "The form will open with all the current information filled in.",
        "Change only the fields you need to update. Leave the rest as they are.",
        "When done, click the blue \"Update\" button at the bottom.",
        "A green message will say \"Record updated successfully\".",
    ])

    add_screenshot_placeholder(doc, "Figure 5.1 - The Edit button on a staff profile")

    add_heading(doc, "What happens when you save", level=2)
    add_para(doc, "When you click \"Update\", three things happen at once:")
    add_steps(doc, [
        "The new information replaces the old information in the system.",
        "The system records who made the change, and the date and time.",
        "A note is added to the staff member's History tab so the change can be traced later.",
    ])

    add_warning_box(doc, "IMPORTANT", [
        "Once you click \"Update\", the change is saved straight away.",
        "There is no \"undo\" button. If you make a mistake, you must edit the record again to correct it.",
        "Always read your changes carefully before clicking \"Update\".",
    ])

    doc.add_page_break()

    # === 6. PROMOTION / TRANSFER ===
    add_heading(doc, "6. How to Record a Promotion or Transfer", level=1)
    add_para(doc,
        "Promotions and transfers are not recorded by editing the main "
        "profile. They have their own special pages so the system can keep "
        "a proper history."
    )

    add_heading(doc, "To record a promotion", level=2)
    add_steps(doc, [
        "Open the staff member's profile.",
        "Click the \"Promotions\" tab.",
        "Click the green \"+ Add Promotion\" button.",
        "Fill in the form: New Grade Level, New Step, Effective Date, and Approval Reference (the letter or memo number).",
        "Attach the promotion letter by clicking \"Upload Document\".",
        "Click \"Save Promotion\".",
        "The staff member's current grade will update automatically.",
    ])

    add_screenshot_placeholder(doc, "Figure 6.1 - The Add Promotion form")

    add_heading(doc, "To record a transfer", level=2)
    add_steps(doc, [
        "Open the staff member's profile.",
        "Click the \"Employment\" tab.",
        "Click the \"Record Transfer\" button.",
        "Fill in the form: New Department, New Designation, Effective Date, and Reason.",
        "Attach the transfer memo by clicking \"Upload Document\".",
        "Click \"Save Transfer\".",
    ])

    add_tip_box(doc, "TIP", [
        "Always attach the official letter or memo. Without it, the record cannot be verified later.",
        "Use the exact effective date written on the letter, not the date you are entering it into the system.",
    ])

    doc.add_page_break()

    # === 7. PROMOTIONS DUE ===
    add_heading(doc, "7. How to View the Promotions Due List", level=1)
    add_para(doc,
        "The system can tell you which officers are due for promotion. "
        "This helps the Court plan ahead and avoid delays."
    )

    add_steps(doc, [
        "From the dashboard, click \"Promotions\" on the left-hand menu.",
        "Choose \"Promotions Due\" from the sub-menu.",
        "A list will open showing every officer who is due for promotion.",
        "By default, it shows officers due within the next 6 months.",
        "To change the period, click the drop-down at the top and choose 3 months, 6 months, or 12 months.",
        "You can sort the list by clicking any column heading (Name, Grade, Due Date).",
        "To see why an officer is listed, click their name.",
    ])

    add_screenshot_placeholder(doc, "Figure 7.1 - The Promotions Due page")

    add_tip_box(doc, "TIP", [
        "The list is updated automatically every night.",
        "You can download this list as a report - see Section 9.",
    ])

    doc.add_page_break()

    # === 8. RETIREMENT MONITOR ===
    add_heading(doc, "8. How to View the Retirement Monitor", level=1)
    add_para(doc,
        "The Retirement Monitor shows officers approaching retirement age "
        "(60 years) or 35 years of service - whichever comes first."
    )

    add_steps(doc, [
        "From the dashboard, click \"Retirement Monitor\" on the left-hand menu.",
        "A list will open showing officers near retirement.",
        "The list is split into three colours:",
        "    RED - officers retiring within 6 months.",
        "    AMBER (orange) - officers retiring within 12 months.",
        "    GREEN - officers retiring within 24 months.",
        "Click any name to see the officer's full profile and exact retirement date.",
    ])

    add_screenshot_placeholder(doc, "Figure 8.1 - The Retirement Monitor with colour-coded list")

    add_tip_box(doc, "TIP", [
        "Use this list every month to prepare retirement letters and benefits early.",
        "If a retirement date looks wrong, check the date of birth and date of first appointment on the staff profile.",
    ])

    doc.add_page_break()

    # === 9. REPORTS ===
    add_heading(doc, "9. How to Generate and Download a Report", level=1)
    add_para(doc,
        "Reports are useful for meetings, audits, and yearly returns. "
        "The system creates them for you in a few clicks."
    )

    add_steps(doc, [
        "From the dashboard, click \"Reports\" on the left-hand menu.",
        "A list of report types will appear. Common ones are: Nominal Roll, Promotions Due, Retirement List, Staff by Department, and Staff by Grade.",
        "Click the report you want.",
        "Choose the options on the right: department, date range, grade level, and so on.",
        "Click \"Generate Report\".",
        "Wait a few seconds. The report will appear on the screen.",
        "To save it, click \"Download as PDF\" or \"Download as Excel\" at the top of the report.",
        "The file will be saved to your computer's Downloads folder.",
        "Open it from there, or attach it to an email.",
    ])

    add_screenshot_placeholder(doc, "Figure 9.1 - The Reports page with download options")

    add_tip_box(doc, "TIP", [
        "PDF is best for printing.",
        "Excel is best when you want to sort or do further work on the figures.",
        "Reports always show the date and time they were generated, so you know they are current.",
    ])

    doc.add_page_break()

    # === 10. PASSWORD ===
    add_heading(doc, "10. What to Do if You Forget Your Password", level=1)
    add_para(doc, "Do not panic. This happens often and is easy to fix.")

    add_heading(doc, "If you registered an email address", level=2)
    add_steps(doc, [
        "Go to the login page.",
        "Click the small link that says \"Forgot Password?\" below the password box.",
        "Type the email address linked to your account.",
        "Click \"Send Reset Link\".",
        "Open your email inbox. You will see a message from the system.",
        "Click the blue \"Reset My Password\" button inside the email.",
        "A new page will open. Type a new password twice.",
        "Click \"Save New Password\".",
        "Go back to the login page and log in with your new password.",
    ])

    add_screenshot_placeholder(doc, "Figure 10.1 - The Forgot Password link on the login page")

    add_heading(doc, "If you did not register an email address", level=2)
    add_steps(doc, [
        "Call or visit the system administrator.",
        "Ask for a password reset.",
        "The administrator will give you a temporary password.",
        "Log in with the temporary password.",
        "The system will ask you to choose a new password straight away.",
        "Type your new password twice and click \"Save\".",
    ])

    add_tip_box(doc, "TIPS FOR A GOOD PASSWORD", [
        "Use at least 8 characters.",
        "Mix letters, numbers, and one symbol (for example: Court#2026).",
        "Do not use your name, your birthday, or the word \"password\".",
        "Do not write your password on paper near your computer.",
    ])

    doc.add_page_break()

    # === 11. COMMON ERRORS ===
    add_heading(doc, "11. Common Errors and What They Mean", level=1)
    add_para(doc,
        "Now and then the system will show a red or yellow message. "
        "Below are the most common ones and what to do."
    )

    err_table = doc.add_table(rows=1, cols=3)
    err_table.style = 'Light Grid Accent 1'
    hdr = err_table.rows[0].cells
    hdr[0].text = "Message on screen"
    hdr[1].text = "What it means"
    hdr[2].text = "What to do"
    for cell in hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    errors = [
        ("Invalid username or password",
         "You typed the wrong username or password.",
         "Check your typing. Make sure Caps Lock is off. Try again. If it fails three times, use \"Forgot Password\"."),
        ("Your account is locked",
         "You typed the wrong password too many times.",
         "Wait 15 minutes and try again, or call the administrator to unlock it."),
        ("Session expired - please log in again",
         "You have been logged out because you were idle for a long time.",
         "This is normal. Log in again to continue."),
        ("Required field is missing",
         "You tried to save a form without filling a field marked with a red star.",
         "Scroll up. Fields with a red border are the ones missing. Fill them and click Save again."),
        ("File number already exists",
         "Another staff record already uses that file number.",
         "Check the file number. If you are sure it is correct, search for it first - the record may already be in the system."),
        ("Invalid date format",
         "The date was typed in the wrong way.",
         "Use the calendar icon to pick the date instead of typing it."),
        ("Phone number must be 11 digits",
         "The phone number is too short or too long.",
         "Re-enter the phone number. Do not include spaces or dashes."),
        ("File too large",
         "The photo or document you tried to upload is too big.",
         "The limit is 2 MB for photos and 5 MB for documents. Ask the IT officer to resize the file."),
        ("You do not have permission",
         "Your user role does not allow this action.",
         "If you believe you should have access, contact the administrator."),
        ("Connection lost - please check your internet",
         "Your computer lost its connection to the network.",
         "Check the network cable or Wi-Fi. Wait a moment, then refresh the page (press F5)."),
        ("Server error - try again later",
         "Something has gone wrong on the system itself.",
         "Wait a few minutes and try again. If the error keeps coming, call the system administrator."),
    ]
    for msg, mean, todo in errors:
        row = err_table.add_row()
        for i, txt in enumerate((msg, mean, todo)):
            p = row.cells[i].paragraphs[0]
            r = p.add_run(txt)
            r.font.size = Pt(10)
            if i == 0:
                r.bold = True

    doc.add_paragraph()

    add_warning_box(doc, "WHEN IN DOUBT", [
        "If you see any message you do not understand, write it down word-for-word.",
        "Take a photograph of the screen if you can.",
        "Then call the system administrator. Do not keep clicking buttons - this can make things worse.",
    ])

    doc.add_page_break()

    # === CLOSING ===
    add_heading(doc, "End of Guide", level=1)
    add_para(doc,
        "You have reached the end of the training guide. With practice, "
        "every task in this book will feel familiar and easy."
    )
    add_para(doc,
        "Remember: the system is here to help you serve the Court better. "
        "Take your time, double-check your entries, and never be afraid to ask for help."
    )

    add_tip_box(doc, "NEED MORE HELP?", [
        "Speak to the System Administrator at the ICT Unit.",
        "Or ask a colleague who has been trained on the system.",
        "Keep this guide in a safe place at your desk.",
    ])

    sign = doc.add_paragraph()
    sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sign_run = sign.add_run("\n\n--- End of Document ---")
    sign_run.italic = True
    sign_run.font.size = Pt(10)
    sign_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    out_path = r"c:\Users\hp\Documents\CCA STAFF\Staff_Biodata_System_User_Training_Guide.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    build_document()
